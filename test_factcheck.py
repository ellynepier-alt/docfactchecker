"""
Automated regression test suite for factcheck_engine.py.

Run with:  pytest test_factcheck.py -v

This encodes every behavior that was manually verified during development, so future
changes can be checked automatically instead of requiring a person to re-run ad hoc
test scripts. Test fixtures are built on the fly (via python-docx/python-pptx/PIL)
rather than shipped as binary files, so this file is self-contained.
"""
import os
import tempfile

import pytest
from docx import Document
from docx.oxml.ns import qn

import factcheck_engine as fe

KB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'doc_guidelines_kb.json')


@pytest.fixture(scope='module')
def kb():
    return fe.load_kb(KB_PATH)


@pytest.fixture
def tmp_txt():
    """Return a function that writes text to a temp .txt file and returns its path."""
    paths = []

    def _make(content, name='test.txt'):
        d = tempfile.mkdtemp()
        path = os.path.join(d, name)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        paths.append(path)
        return path

    yield _make


def flag_kinds(result):
    return [f['kind'] for f in result['flags']]


# ---------------------------------------------------------------------------
# Relevance gate
# ---------------------------------------------------------------------------
class TestRelevanceGate:
    def test_offtopic_document_flagged(self, kb, tmp_txt):
        path = tmp_txt(
            "Industrial robots are increasingly used in manufacturing settings.\n"
            "Modern robotic arms use servo motors and encoders for precise motion planning.\n"
            "Factors such as payload capacity, reach, and individual joint variability affect robot selection.\n"
        )
        result = fe.run_checks(path, kb)
        assert result['coverage'] == []
        assert 'Document relevance' in flag_kinds(result)

    def test_ontopic_document_not_flagged(self, kb, tmp_txt):
        path = tmp_txt("Clinicians should refer patients to a multidisciplinary rehabilitation team.")
        result = fe.run_checks(path, kb)
        assert 'Document relevance' not in flag_kinds(result)

    def test_single_distinctive_term_is_sufficient(self, kb, tmp_txt):
        """A single highly-distinctive term (amantadine) should be enough — this was a real
        false-negative bug where 2 generic hits were required."""
        path = tmp_txt("Amantadine 100 mg twice daily was given to the patient.")
        result = fe.run_checks(path, kb)
        assert 'Document relevance' not in flag_kinds(result)


# ---------------------------------------------------------------------------
# Coverage signature accuracy
# ---------------------------------------------------------------------------
class TestCoverageSignatures:
    def test_every_rec_has_a_signature(self, kb):
        rec_ids = {r['id'] for r in kb['recommendations']}
        assert rec_ids == set(fe.REC_COVERAGE_SIGNATURES.keys())

    def test_every_signature_regex_compiles(self):
        import re
        for patterns in fe.REC_COVERAGE_SIGNATURES.values():
            for p in patterns:
                re.compile(p)  # raises re.error if invalid

    def test_every_rec_text_matches_its_own_signature(self, kb):
        """Sanity check: each recommendation's own guideline text should trigger its own
        coverage signature (otherwise the signature is miscalibrated)."""
        import re
        for rec in kb['recommendations']:
            norm = re.sub(r'\s+', ' ', rec['text'].lower())
            patterns = fe.REC_COVERAGE_SIGNATURES[rec['id']]
            assert any(re.search(p, norm) for p in patterns), f"Rec {rec['id']} text doesn't match its own signature"

    def test_paraphrase_multidisciplinary_team(self, kb, tmp_txt):
        """Regression: 'multidisciplinary team' (not just 'multidisciplinary rehabilitation')
        must count toward Rec 1."""
        path = tmp_txt("Early intervention by a multidisciplinary team can enhance patient outcomes.")
        result = fe.run_checks(path, kb)
        assert '1' in result['coverage']


# ---------------------------------------------------------------------------
# Terminology / contradiction checks
# ---------------------------------------------------------------------------
class TestContradictionChecks:
    def test_permanent_vegetative_state_flagged(self, kb, tmp_txt):
        path = tmp_txt("The patient remains in a permanent vegetative state.")
        result = fe.run_checks(path, kb)
        assert any(f['kind'] == 'Terminology' for f in result['flags'])

    def test_coma_with_eyes_open_flagged(self, kb, tmp_txt):
        path = tmp_txt("The patient is in a coma, with eyes remaining open throughout the exam.")
        result = fe.run_checks(path, kb)
        assert any('coma' in f['matched'].lower() for f in result['flags'])

    def test_vs_uws_with_command_following_flagged(self, kb, tmp_txt):
        path = tmp_txt("The patient remains in a vegetative state but is noted to be following commands consistently.")
        result = fe.run_checks(path, kb)
        assert any('purposeful behavior' in f['matched'] or 'commands' in f['matched'] for f in result['flags'])

    def test_mcs_plus_with_mcs_minus_signs_flagged(self, kb, tmp_txt):
        path = tmp_txt("The patient is classified as MCS+ with automatic movements and object manipulation, but no command following or speech.")
        result = fe.run_checks(path, kb)
        assert any('MCS+' in f['matched'] for f in result['flags'])

    def test_untrained_nonstandardized_flagged(self, kb, tmp_txt):
        path = tmp_txt("Untrained clinicians should use non standardized assessments to assess patients.")
        result = fe.run_checks(path, kb)
        assert any('non-standardized' in f['matched'] for f in result['flags'])

    def test_early_withdrawal_of_life_support_flagged(self, kb, tmp_txt):
        path = tmp_txt("Clinicians are advised to recommend withdrawal of life sustaining treatments before the 72 hour mark.")
        result = fe.run_checks(path, kb)
        assert any('withdrawal' in f['matched'] for f in result['flags'])


# ---------------------------------------------------------------------------
# Numeric key-fact checks
# ---------------------------------------------------------------------------
class TestKeyFactChecks:
    def test_wrong_amantadine_dose_flagged(self, kb, tmp_txt):
        path = tmp_txt("Amantadine 500 mg twice daily was given.")
        result = fe.run_checks(path, kb)
        assert any('amantadine 500 mg' in f['matched'] for f in result['flags'])

    def test_correct_amantadine_dose_not_flagged_as_mismatch(self, kb, tmp_txt):
        path = tmp_txt("Amantadine 100 mg twice daily was given.")
        result = fe.run_checks(path, kb)
        assert not any(f['kind'] == 'Key fact mismatch' and 'amantadine' in f['matched'] for f in result['flags'])

    def test_low_crsr_score_called_favorable_flagged(self, kb, tmp_txt):
        path = tmp_txt("The patient has a CRS-R score of 3, which the team feels is discouraging.")
        result = fe.run_checks(path, kb)
        assert any('CRS-R score' in f['matched'] for f in result['flags'])

    def test_traumatic_3_month_rule_mismatch_flagged(self, kb, tmp_txt):
        path = tmp_txt("Use of the term permanent VS should be discontinued after 3 months in traumatic VS/UWS.")
        result = fe.run_checks(path, kb)
        assert len(result['flags']) > 0


# ---------------------------------------------------------------------------
# Quiz detection and answer verification
# ---------------------------------------------------------------------------
class TestQuizHandling:
    QUESTION = (
        'The term "permanent vegetative state (VS)" should be discontinued after which time '
        'point for patients with traumatic brain injury (TBI)?\n'
        'A) 28 days\nB) 3 months\nC) 6 months\nD) 12 months\n'
    )

    def test_correct_answer_produces_no_flags(self, kb, tmp_txt):
        path = tmp_txt(self.QUESTION + "Answer: D\n")
        result = fe.run_checks(path, kb)
        assert result['flags'] == []

    def test_wrong_answer_is_caught(self, kb, tmp_txt):
        path = tmp_txt(self.QUESTION + "Answer: B\n")
        result = fe.run_checks(path, kb)
        assert any(f['kind'] == 'Possible quiz answer error' for f in result['flags'])

    def test_question_stem_not_flagged_as_assertion(self, kb, tmp_txt):
        """The stem mentions 'permanent vegetative state' as the SUBJECT of the question,
        not as an assertion — it must not trigger the terminology flag."""
        path = tmp_txt(self.QUESTION + "Answer: D\n")
        result = fe.run_checks(path, kb)
        assert not any(f['kind'] == 'Terminology' for f in result['flags'])

    def test_multi_question_isolation(self, kb, tmp_txt):
        """A wrong answer in question 2 must be caught without question 1 or 3
        (both correct) contaminating or being contaminated by it."""
        text = (
            "Question 1: A patient remains in a vegetative state but is noted to be following "
            "commands consistently. What is the correct classification?\n"
            "A) Coma\nB) Vegetative State\nC) Minimally Conscious State\nD) Locked-in Syndrome\nAnswer: C\n\n"
            + self.QUESTION + "Answer: B\n\n"
            "Question 3: Is a coma consistent with eyes remaining open?\nA) Yes\nB) No\nAnswer: B\n"
        )
        path = tmp_txt(text)
        result = fe.run_checks(path, kb)
        assert len(result['flags']) == 1
        assert result['flags'][0]['kind'] == 'Possible quiz answer error'

    def test_crsr_subscale_breakdown_not_falsely_flagged(self, kb, tmp_txt):
        """Regression: a CRS-R SUBSCALE breakdown (Auditory: 3, Visual: 1, ...) must not be
        misread as an overall total-score claim."""
        path = tmp_txt(
            "Based on the CRS-R exam, which subscale scores best represent this presentation?\n"
            "A) Auditory: 3, Visual: 1, Motor: 1, Oromotor/Verbal: 3, Communication: 0, Arousal: 2\n"
            "B) Auditory: 0, Visual: 0, Motor: 0, Oromotor/Verbal: 0, Communication: 0, Arousal: 0\n"
            "Answer: A\n"
        )
        result = fe.run_checks(path, kb)
        assert result['flags'] == []

    def test_genuine_crsr_total_score_error_still_caught(self, kb, tmp_txt):
        path = tmp_txt(
            "What CRS-R total score is associated with increased likelihood of recovery in "
            "nontraumatic post-anoxic VS/UWS?\n"
            "A) A total score of 2\nB) A total score of 4\nC) A total score of 6 or higher\nD) A total score of 1\n"
            "Answer: B) A total score of 4\n"
        )
        result = fe.run_checks(path, kb)
        assert any(f['kind'] == 'Possible quiz answer error' for f in result['flags'])


# ---------------------------------------------------------------------------
# Proofreading
# ---------------------------------------------------------------------------
class TestProofreading:
    def test_abbreviations_not_flagged_as_capitalization_errors(self):
        text = 'The patient was assessed, e.g. clinicians and therapists. Please see Fig. 2 for details.'
        issues = fe.check_grammar(text)
        assert issues == []

    def test_numbered_list_markers_not_flagged(self):
        text = 'Steps: 1. first do this. 2. then do that.'
        issues = fe.check_grammar(text)
        assert issues == []

    def test_genuine_repeated_word_still_caught(self):
        text = 'The the patient was assessed.'
        issues = fe.check_grammar(text)
        assert any('Repeated word' in i['issue'] for i in issues)

    def test_repeated_word_does_not_span_heading_boundary(self):
        """Regression: 'Existing Research\\nResearch on...' across a heading/paragraph
        boundary must not be flagged as a repeated word."""
        text = 'Limitations of Existing Research\nResearch on the use of guided imagery shows promise.'
        issues = fe.check_grammar(text)
        assert issues == []

    def test_references_section_excluded_from_checks(self, kb):
        text = (
            "Conclusion\nThis is the end of the document.\n\n"
            "References\n"
            "Smith J. Recovery outcomes. Psychiatry, 16(5), 340-347. https://doi.org/10.1000/xyz\n"
        )
        body = fe.strip_references_section(text)
        assert 'references' not in body.lower()
        assert 'Conclusion' in body

    def test_curly_and_straight_quotes_both_count_as_attributed(self, kb):
        """Regression: Word's default curly quotes must be recognized, not just straight quotes."""
        rec1_text = kb['recommendations'][0]['text']
        quoted_sentence = f'The guideline states \u201c{rec1_text}\u201d which is directly cited here.'
        findings = fe.check_plagiarism(quoted_sentence, kb)
        if findings:
            assert findings[0]['quoted'] is True


# ---------------------------------------------------------------------------
# Accessibility checks (docx)
# ---------------------------------------------------------------------------
class TestAccessibilityDocx:
    def _make_docx(self, tmp_path_factory):
        d = tempfile.mkdtemp()
        path = os.path.join(d, 'test.docx')
        doc = Document()
        doc.add_heading('Real heading', level=1)
        doc.add_paragraph('Some normal body text.')
        p = doc.add_paragraph('BOLD FAKE HEADING')
        p.runs[0].bold = True
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header A'
        table.cell(0, 1).text = 'Header B'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        doc.save(path)
        return path

    def test_heading_and_table_findings(self):
        path = self._make_docx(None)
        findings = fe.check_accessibility_docx(path)
        heading_finding = next(f for f in findings if f['check'] == 'Heading styles used for structure')
        assert heading_finding['status'] == 'pass'
        fake_heading_finding = next((f for f in findings if f['check'] == 'Bold text used in place of headings'), None)
        assert fake_heading_finding is not None
        table_finding = next(f for f in findings if f['check'] == 'Table header rows')
        assert table_finding['status'] == 'fail'  # no explicit tblHeader flag set


# ---------------------------------------------------------------------------
# OCR warning
# ---------------------------------------------------------------------------
class TestOcrWarning:
    def test_warning_when_engine_unavailable_for_image(self, kb, tmp_txt, monkeypatch):
        monkeypatch.setitem(fe._OCR_STATUS, 'available', False)
        from PIL import Image
        d = tempfile.mkdtemp()
        path = os.path.join(d, 'test.png')
        Image.new('RGB', (100, 50), color='white').save(path)
        result = fe.run_checks(path, kb)
        assert result['ocr_warning'] is not None
        assert 'Tesseract' in result['ocr_warning']

    def test_no_warning_when_engine_available_and_text_found(self, kb, monkeypatch):
        monkeypatch.setitem(fe._OCR_STATUS, 'available', True)
        warning = fe.compute_ocr_warning('test.png', 'Some extracted text here.')
        assert warning is None

    def test_no_warning_for_plain_text_files(self, kb, tmp_txt):
        path = tmp_txt("Plain text file, no images involved.")
        result = fe.run_checks(path, kb)
        assert result['ocr_warning'] is None


# ---------------------------------------------------------------------------
# File format support
# ---------------------------------------------------------------------------
class TestFileFormatSupport:
    def test_supported_extensions_include_images(self):
        assert '.png' in fe.SUPPORTED_EXTS
        assert '.jpg' in fe.SUPPORTED_EXTS
        assert '.jpeg' in fe.SUPPORTED_EXTS

    def test_docx_text_extraction_basic(self):
        """Basic sanity check that docx extraction works and doesn't duplicate ordinary content.
        (A true text-box-duplication regression test would need a fixture with real
        mc:Choice/mc:Fallback XML, which python-docx's high-level API can't construct —
        that specific bug was verified manually against real uploaded files instead.)"""
        path = self._make_docx_with_textbox()
        text = fe.extract_text(path)
        assert text.count('Regular paragraph text.') == 1

    def _make_docx_with_textbox(self):
        d = tempfile.mkdtemp()
        path = os.path.join(d, 'textbox_test.docx')
        doc = Document()
        doc.add_paragraph('Regular paragraph text.')
        doc.save(path)
        return path


# ---------------------------------------------------------------------------
# Whitespace/hyphen tolerance and paraphrase matching
# (regression tests for a real bug: the core terminology_flags/contradiction_flags
# loop used to do a literal substring match against unnormalized text, so a KB
# pattern split across a paragraph/table-cell/PDF-page boundary — or written with a
# hyphen instead of a space — silently failed to match. flexible_pattern() fixes this;
# these tests keep it fixed.)
# ---------------------------------------------------------------------------
class TestWhitespaceAndParaphraseRobustness:
    def test_phrase_split_across_paragraph_boundary_still_flagged(self, kb):
        """A KB phrase ('permanent vegetative state') split across two separate docx
        paragraphs — the way a PDF page break or table-cell boundary would split it —
        must still be flagged. This is the core bug: extract_text() joins chunks with
        '\\n', and the old literal-substring match against unnormalized text missed
        anything split at one of those boundaries."""
        d = tempfile.mkdtemp()
        path = os.path.join(d, 'split.docx')
        doc = Document()
        doc.add_paragraph('The family was told this is now a permanent')
        doc.add_paragraph('vegetative state with little hope of change.')
        doc.save(path)
        result = fe.run_checks(path, kb)
        assert 'Terminology' in flag_kinds(result)

    def test_hyphenated_variant_flagged(self, kb, tmp_txt):
        """'brain dead' (KB pattern, space-separated) must also match 'brain-dead'
        (hyphenated) — real documents are inconsistent about which one they use."""
        path = tmp_txt('The clinical team described the patient as brain-dead on exam.')
        result = fe.run_checks(path, kb)
        assert any('brain' in f['matched'].lower() for f in result['flags'])

    def test_paraphrased_no_chance_of_recovery_flagged(self, kb, tmp_txt):
        """The literal KB pattern 'no chance of recovery' has a regex sibling entry
        that also catches 'no possibility of recovering consciousness' and similar
        rephrasings — real materials rarely use the guideline's exact wording."""
        path = tmp_txt('The physician explained there is no possibility of recovering consciousness in this case.')
        result = fe.run_checks(path, kb)
        assert any(f['kind'] == 'Possible contradiction' for f in result['flags'])

    def test_flexible_pattern_does_not_bridge_unrelated_sentences(self):
        """Sanity check on the fix itself: flexible_pattern's gap regex must only
        bridge immediate whitespace/period/hyphen — not arbitrary intervening text —
        so it doesn't turn into a fuzzy match across genuinely unrelated sentences."""
        import re
        pattern = fe.flexible_pattern('permanent vegetative state')
        unrelated = 'permanent damage was noted. vegetative state assessment protocols were reviewed separately.'
        assert not re.search(pattern, unrelated, re.IGNORECASE)


if __name__ == '__main__':
    import sys
    sys.exit(pytest.main([__file__, '-v']))
